import json
import os
import sys
from pathlib import Path

import docx
import pytest

from main import (
    _smart_replace_in_paragraph,
    _smart_replace_string,
    main,
)


def _run(args: list[str], tmp_path: Path) -> str:
    """Helper: run main() with given CLI args, return stdout."""
    docx_path = tmp_path / "in.docx"
    docx_out = tmp_path / "out.docx"
    resolved = []
    for a in args:
        if a == "{input}":
            resolved.append(str(docx_path))
        elif a == "{output}":
            resolved.append(str(docx_out))
        else:
            resolved.append(a)
    sys.argv = ["docx-editor"] + resolved
    from io import StringIO
    from contextlib import redirect_stdout

    buf = StringIO()
    with redirect_stdout(buf):
        try:
            main()
        except SystemExit:
            pass
    return buf.getvalue().strip()


def _make_doc(
    tmp_path: Path,
    paragraphs: list[str] | None = None,
    tables: list[list[list[str]]] | None = None,
) -> str:
    """Create a .docx and return its path."""
    d = docx.Document()
    for text in paragraphs or []:
        d.add_paragraph(text)
    for rows in tables or []:
        t = d.add_table(rows=len(rows), cols=len(rows[0]))
        for ri, row in enumerate(rows):
            for ci, text in enumerate(row):
                t.cell(ri, ci).text = text
    path = str(tmp_path / "in.docx")
    d.save(path)
    return path


def _read_text(path: str) -> list[str]:
    return [p.text for p in docx.Document(path).paragraphs]


def _read_table_text(path: str) -> list[list[str]]:
    doc = docx.Document(path)
    result = []
    for t in doc.tables:
        for row in t.rows:
            result.append([c.text for c in row.cells])
    return result


# ── list-paragraphs ──────────────────────────────────────────────


def test_list_paragraphs_plain(capsys, tmp_path):
    _make_doc(tmp_path, ["Hello", "World"])
    sys.argv = ["docx-editor", "-i", str(tmp_path / "in.docx"), "list-paragraphs"]
    main()
    out = capsys.readouterr().out
    assert "[0] Hello" in out
    assert "[1] World" in out


def test_list_paragraphs_json(capsys, tmp_path):
    _make_doc(tmp_path, ["Hello", "World"])
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "--json",
        "list-paragraphs",
    ]
    main()
    data = json.loads(capsys.readouterr().out)
    assert data == [{"index": 0, "text": "Hello"}, {"index": 1, "text": "World"}]


# ── list-tables ──────────────────────────────────────────────────


def test_list_tables_plain(capsys, tmp_path):
    _make_doc(tmp_path, tables=[[["A", "B"], ["C", "D"]]])
    sys.argv = ["docx-editor", "-i", str(tmp_path / "in.docx"), "list-tables"]
    main()
    out = capsys.readouterr().out
    assert "Table 0" in out
    assert "1 rows x 2 cols" in out or "2 rows x 2 cols" in out
    assert "A | B" in out
    assert "C | D" in out


def test_list_tables_json(capsys, tmp_path):
    _make_doc(tmp_path, tables=[[["X", "Y"]]])
    sys.argv = ["docx-editor", "-i", str(tmp_path / "in.docx"), "--json", "list-tables"]
    main()
    data = json.loads(capsys.readouterr().out)
    assert len(data) == 1
    assert data[0]["rows"] == 1
    assert data[0]["cols"] == 2
    assert data[0]["cells"] == [["X", "Y"]]


# ── replace (non-smart) ──────────────────────────────────────────


def test_replace_basic(capsys, tmp_path):
    _make_doc(tmp_path, ["Hello {{name}}"])
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "replace",
        "--old",
        "{{name}}",
        "--new",
        "Alice",
    ]
    main()
    assert _read_text(str(tmp_path / "in.docx")) == ["Hello Alice"]
    assert (
        capsys.readouterr().out.strip()
        == "Replaced '{{name}}' with 'Alice' (include_tables=True)"
    )


def test_replace_no_tables(capsys, tmp_path):
    d = docx.Document()
    d.add_paragraph("{{x}}")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "{{x}}"
    d.save(str(tmp_path / "in.docx"))
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "replace",
        "--old",
        "{{x}}",
        "--new",
        "y",
        "--skip-tables",
    ]
    main()
    texts = _read_text(str(tmp_path / "in.docx"))
    assert "y" in texts[0]
    assert _read_table_text(str(tmp_path / "in.docx"))[0][0] == "{{x}}"


# ── replace --smart ──────────────────────────────────────────────


def test_smart_replace_single_run(capsys, tmp_path):
    path = _make_doc(tmp_path, ["Hello {{name}}"])
    sys.argv = [
        "docx-editor",
        "-i",
        path,
        "replace",
        "--old",
        "{{name}}",
        "--new",
        "Alice",
        "--smart",
    ]
    main()
    assert _read_text(path) == ["Hello Alice"]
    assert "smart=True" in capsys.readouterr().out


def test_smart_replace_cross_run(tmp_path):
    d = docx.Document()
    p = d.add_paragraph()
    p.add_run("{{ x is ")
    p.add_run("True }}")
    path = str(tmp_path / "in.docx")
    d.save(path)
    count = _smart_replace_in_paragraph(p, "{{ x is True }}", "{{ x == True }}")
    assert count == 1
    texts = [r.text for r in p.runs]
    assert texts[0] == "{{ x == "
    assert texts[1] == "True }}"


def test_smart_replace_suffix_preservation(tmp_path):
    d = docx.Document()
    p = d.add_paragraph()
    p.add_run("Value is ")
    p.add_run("True")
    run3 = p.add_run(" }}")
    path = str(tmp_path / "in.docx")
    d.save(path)
    count = _smart_replace_in_paragraph(p, "True }}", "False }}")
    assert count == 1
    texts = [r.text for r in p.runs]
    assert texts[0] == "Value is "
    assert texts[1] == "False"
    assert texts[2] == " }}"
    # run3 completely unchanged (same object)
    assert run3.text == " }}"


def test_smart_replace_multiple_occurrences(tmp_path):
    d = docx.Document()
    p = d.add_paragraph()
    p.add_run("a {{x}} b {{x}} c")
    path = str(tmp_path / "in.docx")
    d.save(path)
    count = _smart_replace_in_paragraph(p, "{{x}}", "{{y}}")
    assert count == 2
    assert "".join(r.text for r in p.runs) == "a {{y}} b {{y}} c"


def test_smart_replace_with_tables(capsys, tmp_path):
    _make_doc(tmp_path, ["Hello {{name}}"], [[["Table {{name}}"]]])
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "replace",
        "--old",
        "{{name}}",
        "--new",
        "Alice",
        "--smart",
    ]
    main()
    assert _read_text(str(tmp_path / "in.docx")) == ["Hello Alice"]
    assert _read_table_text(str(tmp_path / "in.docx"))[0][0] == "Table Alice"
    assert "2 instance(s)" in capsys.readouterr().out


def test_smart_replace_no_match(capsys, tmp_path):
    path = _make_doc(tmp_path, ["Hello world"])
    sys.argv = [
        "docx-editor",
        "-i",
        path,
        "replace",
        "--old",
        "{{x}}",
        "--new",
        "y",
        "--smart",
    ]
    main()
    assert "0 instance(s)" in capsys.readouterr().out


# ── replace-up-to ────────────────────────────────────────────────


def test_replace_up_to_basic(capsys, tmp_path):
    _make_doc(tmp_path, ["Line A {{x}}", "Line B {{x}}", "Line C {{x}}"])
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "replace-up-to",
        "--old",
        "{{x}}",
        "--new",
        "y",
        "--paragraph",
        "2",
    ]
    main()
    texts = _read_text(str(tmp_path / "in.docx"))
    assert "y" in texts[0]
    assert "y" in texts[1]
    assert "{{x}}" in texts[2]
    assert "up to paragraph 2" in capsys.readouterr().out


def test_replace_up_to_smart(capsys, tmp_path):
    d = docx.Document()
    d.add_paragraph("A {{x}}")
    d.add_paragraph("B {{x}}")
    d.add_paragraph("C {{x}}")
    path = str(tmp_path / "in.docx")
    d.save(path)
    count = _smart_replace_string(docx.Document(path), "{{x}}", "y", max_paragraph=2)
    assert count == 2
    texts = [p.text for p in docx.Document(path).paragraphs]
    # _smart_replace_string modifies in-memory, doesn't save; load original
    sys.argv = [
        "docx-editor",
        "-i",
        path,
        "replace-up-to",
        "--old",
        "{{x}}",
        "--new",
        "y",
        "--paragraph",
        "2",
        "--smart",
    ]
    main()
    texts = _read_text(path)
    assert "y" in texts[0]
    assert "y" in texts[1]
    assert "{{x}}" in texts[2]
    assert "smart=True" in capsys.readouterr().out


# ── show ──────────────────────────────────────────────────────────


def test_show_found(capsys, tmp_path):
    _make_doc(tmp_path, ["Hello World", "Goodbye"])
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "show",
        "--text",
        "World",
    ]
    main()
    assert "Hello World" in capsys.readouterr().out


def test_show_not_found(capsys, tmp_path):
    _make_doc(tmp_path, ["Hello World"])
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "show",
        "--text",
        "Nope",
    ]
    main()
    assert "No matches" in capsys.readouterr().out


def test_show_json(capsys, tmp_path):
    _make_doc(tmp_path, ["Hello World"])
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "--json",
        "show",
        "--text",
        "World",
    ]
    main()
    assert json.loads(capsys.readouterr().out) == {"found": True, "text": "Hello World"}


def test_show_json_not_found(capsys, tmp_path):
    _make_doc(tmp_path, ["Hello World"])
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "--json",
        "show",
        "--text",
        "Nope",
    ]
    main()
    assert json.loads(capsys.readouterr().out) == {"found": False, "text": ""}


# ── remove-lines ─────────────────────────────────────────────────


def test_remove_lines(capsys, tmp_path):
    _make_doc(tmp_path, ["Line A", "Line B", "Line C", "Line D"])
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "remove-lines",
        "--first-line",
        "Line B",
        "--count",
        "2",
    ]
    main()
    texts = _read_text(str(tmp_path / "in.docx"))
    assert texts == ["Line A", "Line D"]


# ── set-table-cell ────────────────────────────────────────────────


def test_set_table_cell(capsys, tmp_path):
    _make_doc(tmp_path, tables=[[["old"]]])
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "set-table-cell",
        "--table-idx",
        "0",
        "--row",
        "0",
        "--col",
        "0",
        "--text",
        "new",
    ]
    main()
    assert _read_table_text(str(tmp_path / "in.docx"))[0][0] == "new"


# ── set-table-font-size ──────────────────────────────────────────


def test_set_table_font_size(capsys, tmp_path):
    _make_doc(tmp_path, tables=[[["text"]]])
    d = docx.Document(str(tmp_path / "in.docx"))
    run = d.tables[0].cell(0, 0).paragraphs[0].add_run("text")
    run.font.size = docx.shared.Pt(10)
    d.save(str(tmp_path / "in.docx"))
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "set-table-font-size",
        "--table-idx",
        "0",
        "--size",
        "18",
    ]
    main()
    reopened = docx.Document(str(tmp_path / "in.docx"))
    cell_run = reopened.tables[0].cell(0, 0).paragraphs[0].runs[0]
    assert cell_run.font.size == docx.shared.Pt(18)


# ── output flag ──────────────────────────────────────────────────


def test_output_flag(tmp_path):
    _make_doc(tmp_path, ["Hello {{name}}"])
    out_path = str(tmp_path / "out.docx")
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "-o",
        out_path,
        "replace",
        "--old",
        "{{name}}",
        "--new",
        "Alice",
    ]
    from io import StringIO
    from contextlib import redirect_stdout

    with redirect_stdout(StringIO()):
        main()
    assert _read_text(str(tmp_path / "in.docx")) == ["Hello {{name}}"]
    assert _read_text(out_path) == ["Hello Alice"]


def test_backup_created(tmp_path):
    path = _make_doc(tmp_path, ["original"])
    sys.argv = [
        "docx-editor",
        "-i",
        path,
        "replace",
        "--old",
        "original",
        "--new",
        "modified",
    ]
    from io import StringIO
    from contextlib import redirect_stdout

    with redirect_stdout(StringIO()):
        main()
    assert os.path.exists(path + ".bak")
    assert _read_text(path + ".bak") == ["original"]
    assert _read_text(path) == ["modified"]


# ── error cases ──────────────────────────────────────────────────


def test_file_not_found(capsys, tmp_path):
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "nonexistent.docx"),
        "list-paragraphs",
    ]
    with pytest.raises(SystemExit):
        main()
    assert "not found" in capsys.readouterr().err


def test_table_index_out_of_range(capsys, tmp_path):
    _make_doc(tmp_path, [])
    sys.argv = [
        "docx-editor",
        "-i",
        str(tmp_path / "in.docx"),
        "set-table-cell",
        "--table-idx",
        "99",
        "--row",
        "0",
        "--col",
        "0",
        "--text",
        "x",
    ]
    with pytest.raises(SystemExit):
        main()
    assert "out of range" in capsys.readouterr().err


# ── internal: _smart_replace_in_paragraph ────────────────────────


class TestSmartReplaceInParagraph:
    def test_single_run(self):
        d = docx.Document()
        p = d.add_paragraph("Hello {{name}}")
        assert _smart_replace_in_paragraph(p, "{{name}}", "World") == 1
        assert p.text == "Hello World"

    def test_no_match(self):
        d = docx.Document()
        p = d.add_paragraph("Hello")
        assert _smart_replace_in_paragraph(p, "x", "y") == 0
        assert p.text == "Hello"

    def test_cross_run_boundary(self):
        d = docx.Document()
        p = d.add_paragraph()
        p.add_run("{{ x is ")
        p.add_run("True }}")
        assert _smart_replace_in_paragraph(p, "{{ x is True }}", "{{ x == True }}") == 1
        texts = [r.text for r in p.runs]
        assert texts[0] == "{{ x == "
        assert texts[1] == "True }}"

    def test_suffix_preserved(self):
        d = docx.Document()
        p = d.add_paragraph()
        p.add_run("Value is ")
        p.add_run("True")
        p.add_run(" }}")
        assert _smart_replace_in_paragraph(p, "True }}", "False }}") == 1
        texts = [r.text for r in p.runs]
        assert texts == ["Value is ", "False", " }}"]

    def test_suffix_preserved_multi_line(self):
        d = docx.Document()
        p = d.add_paragraph()
        p.add_run("{{ x is True }}")
        # Same text in one run, suffix match irrelevant but should still work
        assert _smart_replace_in_paragraph(p, "True }}", "False }}") == 1
        assert p.text == "{{ x is False }}"

    def test_multiple_matches(self):
        d = docx.Document()
        p = d.add_paragraph("a {{x}} b {{x}} c")
        assert _smart_replace_in_paragraph(p, "{{x}}", "{{y}}") == 2
        assert p.text == "a {{y}} b {{y}} c"

    def test_empty_new_string(self):
        d = docx.Document()
        p = d.add_paragraph("remove {{this}} please")
        assert _smart_replace_in_paragraph(p, "{{this}}", "") == 1
        assert p.text == "remove  please"

    def test_overlap_independent_runs(self):
        d = docx.Document()
        p = d.add_paragraph()
        p.add_run("keep1 ")
        p.add_run("replace_me")
        p.add_run(" keep2")
        assert _smart_replace_in_paragraph(p, "replace_me", "done") == 1
        texts = [r.text for r in p.runs]
        assert texts == ["keep1 ", "done", " keep2"]

    def test_text_inside_hyperlink(self):
        """Regression: text inside hyperlinks should be matched (paragraph.runs misses nested runs)."""
        from lxml import etree as _etree
        from docx.oxml.ns import qn as _qn

        d = docx.Document()
        p = d.add_paragraph()
        p.add_run("before ")
        hyperlink = _etree.SubElement(p._p, _qn("w:hyperlink"))
        hl_run = _etree.SubElement(hyperlink, _qn("w:r"))
        hl_t = _etree.SubElement(hl_run, _qn("w:t"))
        hl_t.text = "{{ hidden }}"
        p.add_run(" after")
        assert _smart_replace_in_paragraph(p, "{{ hidden }}", "revealed") == 1
        assert p.text == "before revealed after"

    def test_text_inside_hyperlink_in_table_cell(self):
        """Regression: table cells with hyperlinked runs should be matchable by smart replace."""
        from lxml import etree as _etree
        from docx.oxml.ns import qn as _qn

        d = docx.Document()
        t = d.add_table(rows=1, cols=1)
        p = t.cell(0, 0).paragraphs[0]
        p.add_run("prefix ")
        hyperlink = _etree.SubElement(p._p, _qn("w:hyperlink"))
        hl_run = _etree.SubElement(hyperlink, _qn("w:r"))
        hl_t = _etree.SubElement(hl_run, _qn("w:t"))
        hl_t.text = "|request %}"
        p.add_run(" suffix")
        count = _smart_replace_string(
            d, "|request %}", "== true %}", include_tables=True
        )
        assert count == 1
        assert t.cell(0, 0).text == "prefix == true %} suffix"


# ── internal: _smart_replace_string ──────────────────────────────


def test_smart_replace_string_paragraphs_only(tmp_path):
    d = docx.Document()
    d.add_paragraph("A {{x}}")
    d.add_paragraph("B {{x}}")
    d.add_paragraph("C")
    count = _smart_replace_string(d, "{{x}}", "y", include_tables=False)
    assert count == 2
    assert [p.text for p in d.paragraphs] == ["A y", "B y", "C"]


def test_smart_replace_string_with_tables(tmp_path):
    d = docx.Document()
    d.add_paragraph("P {{x}}")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "T {{x}}"
    count = _smart_replace_string(d, "{{x}}", "y", include_tables=True)
    assert count == 2


def test_smart_replace_string_max_paragraph(tmp_path):
    d = docx.Document()
    d.add_paragraph("A {{x}}")
    d.add_paragraph("B {{x}}")
    d.add_paragraph("C {{x}}")
    count = _smart_replace_string(d, "{{x}}", "y", max_paragraph=2)
    assert count == 2
    texts = [p.text for p in d.paragraphs]
    assert texts[0] == "A y"
    assert texts[1] == "B y"
    assert texts[2] == "C {{x}}"
