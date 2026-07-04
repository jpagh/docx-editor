import argparse
import json
import logging
import shutil
import sys

import docx
from docx.oxml.ns import qn
import docxedit

logging.basicConfig(level=logging.WARNING)
_DOCXEDIT_LOGGER = logging.getLogger("docxedit")
_DOCXEDIT_LOGGER.setLevel(logging.CRITICAL)


def _save_doc(doc, args):
    if args.output:
        doc.save(args.output)
    else:
        bak = args.input + ".bak"
        shutil.copy2(args.input, bak)
        doc.save(args.input)


def _list_paragraphs(doc, args):
    items = [{"index": i, "text": p.text} for i, p in enumerate(doc.paragraphs)]
    if args.json:
        print(json.dumps(items))
    else:
        for item in items:
            print(f"[{item['index']}] {item['text']}")


def _list_tables(doc, args):
    tables = []
    for i, table in enumerate(doc.tables):
        cells = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append(
            {
                "index": i,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "cells": cells,
            }
        )
    if args.json:
        print(json.dumps(tables))
    else:
        for t in tables:
            print(f"--- Table {t['index']} ({t['rows']} rows x {t['cols']} cols) ---")
            for ri, row in enumerate(t["cells"]):
                print(f"  Row {ri}: {' | '.join(row)}")


_CURLY_QUOTES = str.maketrans(
    {"\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'"}
)


def _norm_quotes(text):
    return text.translate(_CURLY_QUOTES)


def _match_quotes(old, new, orig):
    """Replace straight quotes in *new* with curly quotes from *orig* where *old* had them.

    Matches the N-th occurrence of each quote type independently so that
    "hello" -> "can't" correctly maps both double-quotes.
    """
    result = list(new)
    for qchar in ('"', "'"):
        o_positions = [i for i, c in enumerate(old) if c == qchar]
        n_idx = 0
        for ni, nc in enumerate(new):
            if nc == qchar and n_idx < len(o_positions):
                oi = o_positions[n_idx]
                if oi < len(orig) and orig[oi] in {
                    "\u201c",
                    "\u201d",
                    "\u2018",
                    "\u2019",
                }:
                    result[ni] = orig[oi]
                n_idx += 1
    return "".join(result)


def _run_text(r_el):
    t = r_el.find(qn("w:t"))
    return (t.text or "") if t is not None else ""


def _all_run_elements(paragraph):
    return paragraph._p.findall(".//" + qn("w:r"))


def _paragraph_text(paragraph):
    return "".join(_run_text(el) for el in _all_run_elements(paragraph))


def _smart_replace_in_paragraph(paragraph, old, new, match_quotes=False):
    run_els = _all_run_elements(paragraph)
    full_text = _paragraph_text(paragraph)

    match_old = _norm_quotes(old) if match_quotes else old
    match_text = _norm_quotes(full_text) if match_quotes else full_text

    occurrences = []
    pos = match_text.find(match_old)
    while pos != -1:
        occurrences.append(pos)
        pos = match_text.find(match_old, pos + 1)

    if not occurrences:
        return 0

    for pos in reversed(occurrences):
        match_end = pos + len(old)

        if match_quotes:
            orig_match = full_text[pos : pos + len(old)]
            local_new = _match_quotes(_norm_quotes(old), new, orig_match)
        else:
            local_new = new

        affected_indices = []
        affected_starts = []
        affected_ends = []
        accum = 0
        for i, el in enumerate(run_els):
            run_start = accum
            run_end = accum + len(_run_text(el))
            if pos < run_end and match_end > run_start:
                affected_indices.append(i)
                affected_starts.append(run_start)
                affected_ends.append(run_end)
            accum += len(_run_text(el))

        if not affected_indices:
            continue

        remaining_new = local_new

        for idx in range(len(affected_indices) - 1, -1, -1):
            i = affected_indices[idx]
            el = run_els[i]
            t_el = el.find(qn("w:t"))
            run_start = affected_starts[idx]
            run_end = affected_ends[idx]

            match_start_in_run = max(pos, run_start) - run_start
            match_end_in_run = min(match_end, run_end) - run_start

            is_multi = len(affected_indices) > 1
            is_last_run = idx == len(affected_indices) - 1
            is_first_run = idx == 0

            if is_multi and is_last_run:
                old_suffix = t_el.text[match_start_in_run:match_end_in_run]
                if len(old_suffix) > 0 and remaining_new.endswith(old_suffix):
                    remaining_new = remaining_new[: -len(old_suffix)]
                    continue

            if is_first_run:
                t_el.text = (
                    t_el.text[:match_start_in_run]
                    + remaining_new
                    + t_el.text[match_end_in_run:]
                )
            else:
                t_el.text = (
                    t_el.text[:match_start_in_run] + t_el.text[match_end_in_run:]
                )

    return len(occurrences)


def _smart_replace_string(
    doc, old, new, include_tables=True, max_paragraph=None, match_quotes=False
):
    match_old = _norm_quotes(old) if match_quotes else old
    total = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        if max_paragraph is not None and idx >= max_paragraph:
            break
        text = _paragraph_text(paragraph)
        check = _norm_quotes(text) if match_quotes else text
        if match_old in check:
            total += _smart_replace_in_paragraph(paragraph, old, new, match_quotes)
    if include_tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = _paragraph_text(paragraph)
                        check = _norm_quotes(text) if match_quotes else text
                        if match_old in check:
                            total += _smart_replace_in_paragraph(
                                paragraph, old, new, match_quotes
                            )
    return total


def _replace(doc, args):
    if args.match_quotes and not args.smart:
        print("Error: --match-quotes requires --smart", file=sys.stderr)
        sys.exit(1)
    if args.smart:
        count = _smart_replace_string(
            doc,
            args.old,
            args.new,
            include_tables=args.include_tables,
            match_quotes=args.match_quotes,
        )
    else:
        docxedit.replace_string(
            doc, args.old, args.new, include_tables=args.include_tables
        )
        count = None
    _save_doc(doc, args)
    if count is not None:
        flags = "smart=True"
        if args.match_quotes:
            flags += ", match_quotes=True"
        print(f"Replaced '{args.old}' with '{args.new}' ({count} instance(s), {flags})")
    else:
        print(
            f"Replaced '{args.old}' with '{args.new}' (include_tables={args.include_tables})"
        )


def _replace_up_to(doc, args):
    if args.match_quotes and not args.smart:
        print("Error: --match-quotes requires --smart", file=sys.stderr)
        sys.exit(1)
    if args.smart:
        count = _smart_replace_string(
            doc,
            args.old,
            args.new,
            include_tables=True,
            max_paragraph=args.paragraph,
            match_quotes=args.match_quotes,
        )
        _save_doc(doc, args)
        flags = "smart=True"
        if args.match_quotes:
            flags += ", match_quotes=True"
        print(
            f"Replaced '{args.old}' with '{args.new}' up to paragraph {args.paragraph} "
            f"({count} instance(s), {flags})"
        )
    else:
        docxedit.replace_string_up_to_paragraph(doc, args.old, args.new, args.paragraph)
        _save_doc(doc, args)
        print(
            f"Replaced '{args.old}' with '{args.new}' up to paragraph {args.paragraph}"
        )


def _show(doc, args):
    matches = [p.text for p in doc.paragraphs if args.text in p.text]
    if args.json:
        if matches:
            print(json.dumps({"found": True, "text": " | ".join(matches)}))
        else:
            print(json.dumps({"found": False, "text": ""}))
    else:
        if not matches:
            print(f"No matches found for '{args.text}'")
        else:
            for m in matches:
                print(m)


def _remove_lines(doc, args):
    docxedit.remove_lines(doc, args.first_line, args.count, show_errors=False)
    _save_doc(doc, args)
    print(f"Removed lines starting with '{args.first_line}' ({args.count} line(s))")


def _set_table_cell(doc, args):
    try:
        table = doc.tables[args.table_idx]
    except IndexError:
        print(
            f"Error: Table index {args.table_idx} out of range (have {len(doc.tables)} tables)",
            file=sys.stderr,
        )
        sys.exit(1)
    docxedit.add_text_in_table(table, args.row, args.col, args.text)
    _save_doc(doc, args)
    print(f"Set table {args.table_idx} cell ({args.row}, {args.col}) to '{args.text}'")


def _set_table_font_size(doc, args):
    try:
        table = doc.tables[args.table_idx]
    except IndexError:
        print(
            f"Error: Table index {args.table_idx} out of range (have {len(doc.tables)} tables)",
            file=sys.stderr,
        )
        sys.exit(1)
    docxedit.change_table_font_size(table, args.size)
    _save_doc(doc, args)
    print(f"Set font size in table {args.table_idx} to {args.size}")


def _setup_parser():
    parser = argparse.ArgumentParser(
        description="Edit .docx files from the command line."
    )
    parser.add_argument("-i", "--input", required=True, help="Path to the .docx file")
    parser.add_argument(
        "-o", "--output", help="Output path (default: overwrite input with .bak backup)"
    )
    parser.add_argument(
        "--json", action="store_true", help="Machine-readable JSON output"
    )
    parser.add_argument(
        "--verbose", action="store_true", help="Show library info and debug messages"
    )

    sub = parser.add_subparsers(dest="command", required=True)

    p = sub.add_parser("list-paragraphs", help="List all paragraphs with index numbers")
    p.set_defaults(func=_list_paragraphs)

    p = sub.add_parser(
        "list-tables", help="List all tables with dimensions and cell contents"
    )
    p.set_defaults(func=_list_tables)

    p = sub.add_parser(
        "replace",
        help="Replace old text with new text in paragraphs and optionally tables",
    )
    p.add_argument("--old", required=True, help="Text to find")
    p.add_argument("--new", required=True, help="Replacement text")
    p.add_argument(
        "--skip-tables",
        dest="include_tables",
        action="store_false",
        default=True,
        help="Skip replacement in tables",
    )
    p.add_argument(
        "--smart",
        action="store_true",
        help="Match across run boundaries, preserving original run formatting",
    )
    p.add_argument(
        "--match-quotes",
        action="store_true",
        help="Match curly and straight quotes interchangeably (requires --smart)",
    )
    p.set_defaults(func=_replace)

    p = sub.add_parser(
        "replace-up-to",
        help="Replace old text with new text up to a given paragraph index",
    )
    p.add_argument("--old", required=True, help="Text to find")
    p.add_argument("--new", required=True, help="Replacement text")
    p.add_argument(
        "--paragraph",
        required=True,
        type=int,
        help="Stop at this paragraph index (0-based)",
    )
    p.add_argument(
        "--smart",
        action="store_true",
        help="Match across run boundaries, preserving original run formatting",
    )
    p.add_argument(
        "--match-quotes",
        action="store_true",
        help="Match curly and straight quotes interchangeably (requires --smart)",
    )
    p.set_defaults(func=_replace_up_to)

    p = sub.add_parser("show", help="Show lines containing specific text")
    p.add_argument("--text", required=True, help="Text to search for")
    p.set_defaults(func=_show)

    p = sub.add_parser(
        "remove-lines", help="Remove lines starting from first matching text"
    )
    p.add_argument(
        "--first-line", required=True, help="Text in the first line to remove"
    )
    p.add_argument("--count", required=True, type=int, help="Number of lines to remove")
    p.set_defaults(func=_remove_lines)

    p = sub.add_parser("set-table-cell", help="Set the text of a specific table cell")
    p.add_argument("--table-idx", required=True, type=int, help="Table index (0-based)")
    p.add_argument("--row", required=True, type=int, help="Row index (0-based)")
    p.add_argument("--col", required=True, type=int, help="Column index (0-based)")
    p.add_argument("--text", required=True, help="Text to set")
    p.set_defaults(func=_set_table_cell)

    p = sub.add_parser(
        "set-table-font-size", help="Change the font size of all text in a table"
    )
    p.add_argument("--table-idx", required=True, type=int, help="Table index (0-based)")
    p.add_argument("--size", required=True, type=int, help="Font size in points")
    p.set_defaults(func=_set_table_font_size)

    return parser


def main():
    parser = _setup_parser()
    args = parser.parse_args()

    if args.verbose:
        logging.basicConfig(level=logging.INFO, force=True)
        _DOCXEDIT_LOGGER.setLevel(logging.INFO)

    try:
        doc = docx.Document(args.input)
    except FileNotFoundError:
        print(f"Error: File '{args.input}' not found", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error: Could not load document '{args.input}': {e}", file=sys.stderr)
        sys.exit(1)

    args.func(doc, args)


if __name__ == "__main__":
    main()
